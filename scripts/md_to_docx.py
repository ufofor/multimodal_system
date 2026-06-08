"""Convert PM markdown docs to formatted .docx files."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_doc(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)  # EY navy
        h.font.size = Pt(16 - (i * 2))
        h.font.bold = True


def add_table_from_md(doc, lines):
    rows = [l for l in lines if l.startswith("|")]
    if len(rows) < 2:
        return

    # Filter separator row
    data_rows = [r for r in rows if not re.match(r"^\|[\s\-|]+\|$", r)]
    if not data_rows:
        return

    cols = [c.strip() for c in data_rows[0].strip("|").split("|")]
    n_cols = len(cols)
    n_rows = len(data_rows)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    for ri, row_text in enumerate(data_rows):
        cells = [c.strip() for c in row_text.strip("|").split("|")]
        for ci, cell_text in enumerate(cells[:n_cols]):
            cell = table.cell(ri, ci)
            # Strip bold markdown
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            # Strip inline code
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            cell.text = clean
            if ri == 0:
                set_cell_bg(cell, "1F3564")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
            elif ri % 2 == 0:
                set_cell_bg(cell, "EEF2F7")

    doc.add_paragraph()


def parse_md_to_docx(md_path: Path, out_path: Path):
    doc = Document()
    style_doc(doc)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    lines = md_path.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)

        # Table block
        elif line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            add_table_from_md(doc, block)
            continue

        # Code block — skip fence markers, preserve content
        elif line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # Bullet
        elif line.startswith("- "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line[2:])
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")

        # Blockquote
        elif line.startswith("> "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line[2:])
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Bold metadata lines (e.g. **Status:** ...)
        elif line.startswith("**") and "**" in line[2:]:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Normal paragraph
        elif line.strip():
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text)

        i += 1

    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    docs_dir = Path(__file__).parent.parent / "docs"
    for md_file in ["PRD.md", "ADR.md", "METRICS.md", "TEARDOWN.md"]:
        md_path = docs_dir / md_file
        out_path = docs_dir / md_file.replace(".md", ".docx")
        parse_md_to_docx(md_path, out_path)
